"""
Q-TRON Gen3 전략 기준표 DOCX 생성기
python make_strategy_doc.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from pathlib import Path

OUT = Path(__file__).parent / "Gen3_전략기준표.docx"

# ── 색상 ──────────────────────────────────────────────────────────────────
C_HEADER  = RGBColor(0x1F, 0x4E, 0x79)   # 헤더행 배경
C_SUBHDR  = RGBColor(0x2E, 0x75, 0xB6)   # 서브헤더
C_ALT     = RGBColor(0xEB, 0xF3, 0xFB)   # 짝수행 연한 파랑
C_GREEN   = RGBColor(0xE2, 0xEF, 0xDA)
C_YELLOW  = RGBColor(0xFF, 0xF2, 0xCC)
C_ORANGE  = RGBColor(0xFC, 0xE4, 0xD6)
C_BLUE2   = RGBColor(0xDA, 0xEE, 0xF3)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY    = RGBColor(0x88, 0x88, 0x88)
C_RED     = RGBColor(0xC0, 0x00, 0x00)
C_BODY    = RGBColor(0x1A, 0x1A, 0x1A)
C_TITLE   = RGBColor(0x1F, 0x4E, 0x79)
C_SUB     = RGBColor(0x2E, 0x75, 0xB6)


def rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(color))
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "AAAAAA")
        tcBorders.append(b)
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=90, bottom=90, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def cell_text(cell, text, bold=False, color=None, center=False, size=10):
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return cell


def header_row(table, texts, col_widths_cm, bg=None):
    row = table.add_row()
    bg = bg or C_HEADER
    for i, (text, w) in enumerate(zip(texts, col_widths_cm)):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell_text(cell, text, bold=True, color=C_WHITE, center=True, size=10)
    return row


def data_row(table, cells_data, col_widths_cm):
    """
    cells_data: list of (text, bg, bold, center, color)
    """
    row = table.add_row()
    for i, spec in enumerate(cells_data):
        if isinstance(spec, str):
            text, bg, bold, center, color = spec, C_WHITE, False, False, C_BODY
        else:
            text = spec[0]
            bg    = spec[1] if len(spec) > 1 else C_WHITE
            bold  = spec[2] if len(spec) > 2 else False
            center= spec[3] if len(spec) > 3 else False
            color = spec[4] if len(spec) > 4 else C_BODY
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell_text(cell, text, bold=bold, center=center, color=color, size=10)
    return row


def make_table(doc, col_widths_cm):
    n = len(col_widths_cm)
    table = doc.add_table(rows=0, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, w in enumerate(col_widths_cm):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.name = "맑은 고딕"
    p.runs[0].font.color.rgb = C_TITLE if level == 1 else C_SUB
    p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"\u203B {text}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = C_GRAY
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


# ── 문서 생성 ─────────────────────────────────────────────────────────────
doc = Document()

# 페이지 설정 (A4)
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.0)
section.top_margin  = section.bottom_margin = Cm(2.0)

# 기본 스타일
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

# ── 표지 ─────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Q-TRON Gen3")
run.font.name = "맑은 고딕"; run.font.size = Pt(28); run.font.bold = True
run.font.color.rgb = C_TITLE
run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("전략 기준표")
run.font.name = "맑은 고딕"; run.font.size = Pt(22); run.font.bold = True
run.font.color.rgb = C_SUB
run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Strategy Reference Guide  —  Gen03-02  /  2026-03-09")
run.font.name = "맑은 고딕"; run.font.size = Pt(11); run.font.italic = True
run.font.color.rgb = C_GRAY
run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

doc.add_paragraph()

# 구분선
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:color"), "2E75B6")
pBdr.append(bottom); pPr.append(pBdr)

doc.add_paragraph()
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 1. 레짐 판단
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. 레짐 판단 (RegimeDetector)", 1)

W = [4.5, 12.5]
t = make_table(doc, W)
header_row(t, ["항목", "기준"], W)
ALT = C_ALT

rows = [
    [("분석 대상", ALT), ("KOSPI + KOSDAQ 지수 평균", ALT)],
    [("MA 기준선", C_WHITE), ("MA200  (REGIME_MA = 200)", C_WHITE)],
    [("① MA 정배열", ALT), ("ma5 > ma20 > ma60", ALT)],
    [("② 거래량 추세", C_WHITE), ("5일 평균 거래량 > 20일 평균 거래량", C_WHITE)],
    [("③ 모멘텀", ALT), ("20일 수익률 > 0", ALT)],
    [("④ 변동성 안정", C_WHITE), ("ATR(14) <= 60일 평균 ATR", C_WHITE)],
    [("BULL", C_BLUE2, True, False, C_BODY), ("KOSPI/KOSDAQ 평균 점수 >= 2.5", C_BLUE2)],
    [("SIDEWAYS", C_YELLOW, True, False, C_BODY), ("1.5 < 점수 < 2.5", C_YELLOW)],
    [("BEAR", C_ORANGE, True, False, C_RED), ("KOSPI/KOSDAQ 평균 점수 <= 1.5", C_ORANGE)],
]
for rd in rows:
    data_row(t, rd, W)

add_note(doc, "점수 항목 ①~④는 각 1점, KOSPI+KOSDAQ 평균 합계 0~4점으로 레짐 결정")
doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# 2. 유니버스 필터
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. 유니버스 필터 (UniverseBuilder)", 1)

W = [5.0, 12.0]
t = make_table(doc, W)
header_row(t, ["조건", "기준값"], W)

rows = [
    [("대상 시장", ALT), ("KOSPI + KOSDAQ 전체", ALT)],
    [("우선주 제외", C_WHITE), ("종목코드 끝자리 5~9", C_WHITE)],
    [("관리/거래정지 제외", ALT), ("종목명에 '관리' / '거래정지' / '정리매매' 포함 종목", ALT)],
    [("최소 주가", C_WHITE), ("1,000원 이상", C_WHITE)],
    [("최소 시총", ALT), ("1,000억원 이상", ALT)],
    [("최소 일 거래대금", C_WHITE), ("20억원 이상  (5일 평균, 거래량 x 종가)", C_WHITE)],
]
for rd in rows:
    data_row(t, rd, W)

doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# 3. Q-Score 가중치
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Q-Score 가중치 (QScoreEngine)", 1)
add_heading(doc, "3-1. 레짐별 가중치", 2)

W = [4.0, 3.5, 4.0, 5.5]  # 총 17.0
t = make_table(doc, W)
header_row(t, ["서브스코어", "BULL", "SIDEWAYS", "BEAR"], W)

rows = [
    [("Technical", ALT, True), ("50%", ALT, True, True, C_RED), ("30%", ALT, False, True), ("25%", ALT, False, True)],
    [("Demand", C_WHITE, True), ("25%", C_WHITE, False, True), ("25%", C_WHITE, False, True), ("40%", C_WHITE, True, True, C_RED)],
    [("Price", ALT, True), ("15%", ALT, False, True), ("30%", ALT, True, True, C_RED), ("20%", ALT, False, True)],
    [("Alpha", C_WHITE, True), ("10%", C_WHITE, False, True), ("15%", C_WHITE, False, True), ("15%", C_WHITE, False, True)],
]
for rd in rows:
    data_row(t, rd, W)

doc.add_paragraph()
add_heading(doc, "3-2. 서브스코어 계산 기준", 2)

W = [4.5, 12.5]
t = make_table(doc, W)
header_row(t, ["서브스코어", "계산 기준"], W)

rows = [
    [("Technical  (0~1)", ALT, True), ("MA 정배열(ma5>ma20): +0.4  /  20일 모멘텀>0: +0.4  /  거래량 추세: +0.2", ALT)],
    [("Demand  (0~1)", C_WHITE, True), ("(외인순매수 + 기관순매수) / 총거래량  →  0.5 + net_ratio x 5", C_WHITE)],
    [("Price  (0~1)", ALT, True), ("현재가 / 52주 신고가  (신고가 근접도)", ALT)],
    [("Alpha", C_WHITE, True), ("미구현 (0.0 고정) — 공매도비율 / 뉴스 센티먼트 예정", C_WHITE, False, False, C_GRAY)],
]
for rd in rows:
    data_row(t, rd, W)

add_note(doc, "Q-Score 범위: 0.0 ~ 1.0  /  signals_YYYYMMDD.csv에 저장")
doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# 4. TP/SL 설정 기준
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. TP / SL 설정 기준", 1)

W = [5.0, 12.0]
t = make_table(doc, W)
header_row(t, ["항목", "기준"], W)

rows = [
    [("SL 계산", ALT), ("진입가  -  ATR(14)  x  SL배수", ALT)],
    [("TP 계산", C_WHITE), ("진입가  +  (진입가 - SL)  x  2.0     (R:R = 2 : 1)", C_WHITE)],
    [("SL배수 — BULL", C_BLUE2, True), ("4.0   (넓은 손절 / 추세 추종)", C_BLUE2)],
    [("SL배수 — SIDEWAYS", C_YELLOW, True), ("2.5   (중간)", C_YELLOW)],
    [("SL배수 — BEAR", C_ORANGE, True), ("1.0   (좁은 손절 / 보수적)", C_ORANGE)],
]
for rd in rows:
    data_row(t, rd, W)

doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# 5. 진입 전략 (Stage A / B)
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. 진입 전략 (Stage A / B)", 1)

W = [3.5, 6.5, 7.0]
t = make_table(doc, W)
header_row(t, ["구분", "Stage A — Early Entry", "Stage B — Main Strategy"], W)

rows = [
    [("레짐 조건", ALT), ("BULL 전용", ALT, True), ("BULL / SIDEWAYS / BEAR 전체", ALT)],
    [("신호 기준", C_WHITE), ("signals.csv  stage=A 태그\n(없으면 Q-Score 상위 20% 폴백)", C_WHITE), ("signals.csv  stage=B\n(나머지 전체)", C_WHITE)],
    [("추가 조건", ALT), ("갭업 < 3%  (현재가 <= 전일종가 x 1.03)", ALT), ("섹터 노출도 사전 필터 (90% 미만)", ALT)],
    [("타이밍", C_WHITE), ("장 시작 시가 우선", C_WHITE), ("장 중 순차 진입", C_WHITE)],
    [("선진입 전략", ALT), ("추후 별도 전략 (현재 미적용)", ALT, False, False, C_GRAY), ("—", ALT, False, True, C_GRAY)],
]
for rd in rows:
    data_row(t, rd, W)

doc.add_paragraph()
add_heading(doc, "5-1. Stage 태깅 기준 (signal_generator)", 2)

W = [8.0, 9.0]
t = make_table(doc, W)
header_row(t, ["조건", "stage 태그"], W)
rows = [
    [("BULL 레짐 + Q-Score 상위 20%", ALT), ("stage = A", ALT, True)],
    [("그 외 전체", C_WHITE), ("stage = B", C_WHITE, True)],
]
for rd in rows:
    data_row(t, rd, W)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# 6. 포지션 사이징
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. 포지션 사이징 (Stage C)", 1)

W = [5.0, 12.0]
t = make_table(doc, W)
header_row(t, ["항목", "기준"], W)
rows = [
    [("분산 방식", ALT, True), ("균등 분산", ALT, True)],
    [("1종목 투자금", C_WHITE), ("MIN(현금, 총평가금액)  /  MAX_POSITIONS (20)", C_WHITE)],
    [("최대 보유 종목", ALT), ("20개", ALT)],
    [("1종목 최대 비중", C_WHITE), ("20%", C_WHITE)],
]
for rd in rows:
    data_row(t, rd, W)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# 7. 6중 진입 게이트
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. 6중 진입 게이트 (can_enter)", 1)

W = [2.0, 7.5, 7.5]
t = make_table(doc, W)
header_row(t, ["순서", "조건", "한도"], W)
rows = [
    [("①", ALT, False, True), ("일일 손실 한도", ALT), ("-2% 초과 시 신규 진입 차단", ALT)],
    [("②", C_WHITE, False, True), ("월간 DD 한도", C_WHITE), ("-7% 초과 시 신규 진입 차단", C_WHITE)],
    [("③", ALT, False, True), ("최대 보유 종목 수", ALT), ("20개 초과 시 차단", ALT)],
    [("④", C_WHITE, False, True), ("종목당 최대 비중", C_WHITE), ("20% 초과 시 차단", C_WHITE)],
    [("⑤", ALT, False, True), ("섹터 노출 한도", ALT), ("30% 초과 시 차단", ALT)],
    [("⑥", C_WHITE, False, True), ("총 노출도 한도", C_WHITE), ("60% 초과 시 차단", C_WHITE)],
]
for rd in rows:
    data_row(t, rd, W)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# 8. 청산 조건
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. 청산 조건 (ExitLogic)", 1)

W = [3.0, 4.0, 10.0]
t = make_table(doc, W)
header_row(t, ["우선순위", "조건", "기준"], W)
rows = [
    [("1 (최우선)", C_ORANGE, True, True, C_RED), ("SL", C_ORANGE, True, True, C_RED), ("현재가 <= SL", C_ORANGE)],
    [("2", C_WHITE, False, True), ("MAX_HOLD", C_WHITE, True, True), ("보유일 >= 60일", C_WHITE)],
    [("3", ALT, False, True), ("MA20", ALT, True, True), ("현재가 < MA20 (20일 이동평균)", ALT)],
    [("4", C_WHITE, False, True), ("TP", C_WHITE, True, True), ("현재가 >= TP", C_WHITE)],
]
for rd in rows:
    data_row(t, rd, W)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# 9. 리스크 모드
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9. 리스크 모드", 1)

W = [3.5, 6.5, 7.0]
t = make_table(doc, W)
header_row(t, ["모드", "조건", "대응"], W)
rows = [
    [("NORMAL", C_GREEN, True), ("일손실 > -2%  AND  월DD > -7%", C_GREEN), ("정상 진입 허용", C_GREEN)],
    [("SOFT_STOP", C_YELLOW, True), ("일손실 <= -2%", C_YELLOW), ("신규 진입 차단  (기존 포지션 유지)", C_YELLOW)],
    [("HARD_STOP", C_ORANGE, True, False, C_RED), ("월DD <= -7%", C_ORANGE), ("신규 차단 + 손실 순 강제 청산", C_ORANGE)],
]
for rd in rows:
    data_row(t, rd, W)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# 10. 비용 구조
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "10. 비용 구조", 1)

W = [9.0, 8.0]
t = make_table(doc, W)
header_row(t, ["항목", "비율"], W)
rows = [
    [("거래 수수료  (FEE)", ALT), ("0.015%", ALT, False, True)],
    [("슬리피지  (SLIPPAGE)", C_WHITE), ("0.100%", C_WHITE, False, True)],
    [("거래세  (TAX, 매도 시만 적용)", ALT), ("0.180%", ALT, False, True)],
    [("진입 총비용  (FEE + SLIPPAGE)", C_WHITE, True), ("0.115%", C_WHITE, True, True)],
    [("청산 총비용  (FEE + SLIPPAGE + TAX)", ALT, True), ("0.295%", ALT, True, True)],
]
for rd in rows:
    data_row(t, rd, W)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# 11. 미구현 / 추후 전략
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "11. 미구현 / 추후 전략", 1)

W = [9.0, 8.0]
t = make_table(doc, W)
header_row(t, ["항목", "상태"], W)
rows = [
    [("Alpha 스코어  (공매도비율, 뉴스 센티먼트)", ALT), ("미구현  (0.0 고정)", ALT, False, False, C_GRAY)],
    [("선진입 전략  (Stage A 갭업 전략 확장)", C_WHITE), ("추후 별도 전략으로 분리", C_WHITE, False, False, C_GRAY)],
    [("get_investor_trend() 실데이터", ALT), ("opt10059 TR 등 추후 구현", ALT, False, False, C_GRAY)],
    [("opt20006 지수 TR  (DISABLE_INDEX_TR)", C_WHITE), ("비활성화 상태", C_WHITE, False, False, C_GRAY)],
    [("Kiwoom LIVE 실거래  (_send_to_kiwoom)", ALT), ("미구현", ALT, False, False, C_GRAY)],
    [("Gen2 kiwoom_provider.py TR 버그 수정", C_WHITE), ("추후  (Gen3 안정화 후)", C_WHITE, False, False, C_GRAY)],
    [("비정상 종목코드 필터  (숫자 6자리 검증)", ALT), ("UniverseBuilder 추가 예정", ALT, False, False, C_GRAY)],
]
for rd in rows:
    data_row(t, rd, W)

doc.add_paragraph()
add_note(doc, "이 문서는 Gen03-02 코드베이스 실제 구현을 기반으로 자동 생성되었습니다.")

# ── 저장 ─────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"OK: {OUT}")
